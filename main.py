# main.py
import asyncio
import io
import os
import re
import shutil
import tempfile
import time
import uuid
import zipfile
from xml.etree import ElementTree as ET

import json

import mammoth
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from converter.doc_converter import is_doc, to_doc, to_docx
from formatter.applier import apply_format_to_paragraph, apply_page_margins, center_tables_and_images
from formatter.preview import render_paragraph_html
from formatter.template_loader import list_templates, load_template
from parser.context import ChapterContext
from parser.rules import classify_paragraph

EMU_PER_CM = 360000
PT_PER_CM = 72 / 2.54
_WP_DRAWING_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
_VML_NS = "{urn:schemas-microsoft-com:vml}"
_MC_NS = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_VML_STYLE_WIDTH_RE = re.compile(r"(?:^|;)\s*width\s*:\s*([\d.]+)\s*([a-z%]*)", re.I)


def _vml_shape_width_cm(shape) -> float | None:
    if shape.find(f"{_VML_NS}imagedata") is None:
        return None
    style = shape.get("style") or ""
    m = _VML_STYLE_WIDTH_RE.search(style)
    if not m:
        return None
    try:
        val = float(m.group(1))
    except ValueError:
        return None
    unit = (m.group(2) or "pt").lower()
    if unit in ("pt", ""):
        return val / PT_PER_CM
    if unit == "in":
        return val * 2.54
    if unit == "cm":
        return val
    if unit == "mm":
        return val / 10
    if unit == "px":
        return val / 96 * 2.54
    return None


def _extract_image_widths_cm(docx_bytes: bytes) -> list[float]:
    widths = []
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            with zf.open("word/document.xml") as f:
                tree = ET.parse(f)
        skip = set()
        for fb in tree.iter(f"{_MC_NS}Fallback"):
            for d in fb.iter():
                skip.add(id(d))
        for el in tree.iter():
            if id(el) in skip:
                continue
            if el.tag == f"{_WP_DRAWING_NS}extent":
                cx = el.get("cx")
                if cx and cx.isdigit():
                    widths.append(int(cx) / EMU_PER_CM)
            elif el.tag == f"{_VML_NS}shape":
                w = _vml_shape_width_cm(el)
                if w is not None:
                    widths.append(w)
    except (KeyError, ET.ParseError, zipfile.BadZipFile):
        pass
    return widths


def _inject_image_widths(html: str, widths_cm: list[float]) -> str:
    it = iter(widths_cm)

    def repl(m):
        w = next(it, None)
        if w is None:
            return m.group(0)
        return f'<img style="width:{w:.2f}cm;height:auto;max-width:100%" '

    return re.sub(r'<img(?=\s)', repl, html)


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _extract_media_anchors(docx_bytes: bytes) -> list[dict]:
    """Walk top-level body children, return ordered media anchors.

    Each anchor: {id, kind: 'image'|'table', after_para_index}.
    after_para_index = -1 means before the first paragraph.
    Image inside a paragraph anchors to that paragraph's index.
    """
    anchors: list[dict] = []
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            with zf.open("word/document.xml") as f:
                tree = ET.parse(f)
    except (KeyError, ET.ParseError, zipfile.BadZipFile):
        return anchors

    body = tree.getroot().find(f"{_W_NS}body")
    if body is None:
        return anchors

    skip_ids = set()
    for fb in body.iter(f"{_MC_NS}Fallback"):
        for d in fb.iter():
            skip_ids.add(id(d))

    para_idx = -1
    media_counter = 0
    for child in list(body):
        if child.tag == f"{_W_NS}p":
            para_idx += 1
            for el in child.iter():
                if id(el) in skip_ids:
                    continue
                if el.tag in (f"{_WP_DRAWING_NS}extent", f"{_VML_NS}shape"):
                    if el.tag == f"{_VML_NS}shape" and el.find(f"{_VML_NS}imagedata") is None:
                        continue
                    anchors.append({
                        "id": f"media-{media_counter}",
                        "kind": "image",
                        "after_para_index": para_idx,
                    })
                    media_counter += 1
        elif child.tag == f"{_W_NS}tbl":
            anchors.append({
                "id": f"media-{media_counter}",
                "kind": "table",
                "after_para_index": para_idx,
            })
            media_counter += 1
    return anchors


def _inject_media_ids(html: str, anchors: list[dict]) -> str:
    """Assign sequential id="media-N" to each <img> and <table> in HTML order."""
    counter = [0]
    kinds = [a["kind"] for a in anchors]

    def repl(m):
        tag = m.group(1).lower()
        i = counter[0]
        if i >= len(kinds):
            return m.group(0)
        expected = "image" if tag == "img" else "table"
        if kinds[i] != expected:
            return m.group(0)
        counter[0] += 1
        return f'<{m.group(1)} id="media-{i}"'

    return re.sub(r'<(img|table)(?=[\s>])', repl, html, flags=re.I)


_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(_BASE_DIR, "templates")
UPLOADS_DIR = os.path.join(_BASE_DIR, "uploads")
SESSION_TTL_SECONDS = 2 * 3600
CLEANUP_INTERVAL_SECONDS = 30 * 60

os.makedirs(UPLOADS_DIR, exist_ok=True)

app = FastAPI()
app.mount("/static", StaticFiles(directory=os.path.join(_BASE_DIR, "frontend")), name="static")


@app.get("/")
def index():
    return FileResponse(os.path.join(_BASE_DIR, "frontend", "index.html"))


def _session_dir(doc_id: str) -> str:
    safe = re.fullmatch(r"[0-9a-fA-F-]{8,64}", doc_id or "")
    if not safe:
        raise HTTPException(status_code=400, detail="invalid doc_id")
    return os.path.join(UPLOADS_DIR, doc_id)


def _touch(path: str) -> None:
    try:
        os.utime(path, None)
    except OSError:
        pass


def _cleanup_expired_sessions() -> None:
    now = time.time()
    if not os.path.isdir(UPLOADS_DIR):
        return
    for name in os.listdir(UPLOADS_DIR):
        d = os.path.join(UPLOADS_DIR, name)
        if not os.path.isdir(d):
            continue
        try:
            age = now - os.path.getmtime(d)
        except OSError:
            continue
        if age > SESSION_TTL_SECONDS:
            shutil.rmtree(d, ignore_errors=True)


async def _cleanup_loop() -> None:
    while True:
        try:
            _cleanup_expired_sessions()
        except Exception:
            pass
        await asyncio.sleep(CLEANUP_INTERVAL_SECONDS)


@app.on_event("startup")
async def _start_cleanup() -> None:
    _cleanup_expired_sessions()
    asyncio.create_task(_cleanup_loop())


# ─── OMML → LaTeX 转换器 ────────────────────────────────────────────────────
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W_NS_MAIN = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _m(tag: str) -> str:
    return f"{{{_M_NS}}}{tag}"

# OMML 运算符/符号 → LaTeX 映射
_OMML_SYM_MAP = {
    "∑": r"\sum", "∏": r"\prod", "∫": r"\int", "∬": r"\iint",
    "∭": r"\iiint", "∮": r"\oint", "√": r"\sqrt", "∞": r"\infty",
    "±": r"\pm", "∓": r"\mp", "×": r"\times", "÷": r"\div",
    "≤": r"\leq", "≥": r"\geq", "≠": r"\neq", "≈": r"\approx",
    "∈": r"\in", "∉": r"\notin", "⊂": r"\subset", "⊃": r"\supset",
    "⊆": r"\subseteq", "⊇": r"\supseteq", "∪": r"\cup", "∩": r"\cap",
    "→": r"\to", "←": r"\leftarrow", "⇒": r"\Rightarrow", "⇐": r"\Leftarrow",
    "⇔": r"\Leftrightarrow", "↔": r"\leftrightarrow",
    "∀": r"\forall", "∃": r"\exists", "¬": r"\neg",
    "∧": r"\wedge", "∨": r"\vee", "⊕": r"\oplus", "⊗": r"\otimes",
    "α": r"\alpha", "β": r"\beta", "γ": r"\gamma", "δ": r"\delta",
    "ε": r"\varepsilon", "ζ": r"\zeta", "η": r"\eta", "θ": r"\theta",
    "ι": r"\iota", "κ": r"\kappa", "λ": r"\lambda", "μ": r"\mu",
    "ν": r"\nu", "ξ": r"\xi", "π": r"\pi", "ρ": r"\rho",
    "σ": r"\sigma", "τ": r"\tau", "υ": r"\upsilon", "φ": r"\varphi",
    "χ": r"\chi", "ψ": r"\psi", "ω": r"\omega",
    "Α": "A", "Β": "B", "Γ": r"\Gamma", "Δ": r"\Delta",
    "Ε": "E", "Ζ": "Z", "Η": "H", "Θ": r"\Theta",
    "Λ": r"\Lambda", "Μ": "M", "Ν": "N", "Ξ": r"\Xi",
    "Π": r"\Pi", "Ρ": "R", "Σ": r"\Sigma", "Τ": "T",
    "Υ": r"\Upsilon", "Φ": r"\Phi", "Χ": "X", "Ψ": r"\Psi",
    "Ω": r"\Omega",
    "′": "'", "″": "''", "∂": r"\partial", "∇": r"\nabla",
    "·": r"\cdot", "…": r"\ldots", "⋯": r"\cdots", "⋮": r"\vdots",
    "⋱": r"\ddots", "‖": r"\|", "∥": r"\|",
    "⟨": r"\langle", "⟩": r"\rangle", "⌈": r"\lceil", "⌉": r"\rceil",
    "⌊": r"\lfloor", "⌋": r"\rfloor",
    "ℝ": r"\mathbb{R}", "ℤ": r"\mathbb{Z}", "ℕ": r"\mathbb{N}",
    "ℚ": r"\mathbb{Q}", "ℂ": r"\mathbb{C}",
    "\u2212": "-",  # MINUS SIGN → 普通减号
    "\u2061": "",   # FUNCTION APPLICATION (不可见)
    "\u2062": "",   # INVISIBLE TIMES
    "\u2063": ",",  # INVISIBLE SEPARATOR
    "\u2064": "",   # INVISIBLE PLUS
    "\u2219": r"\bullet",  # BULLET OPERATOR
    "\u2220": r"\angle",
    "\u22c5": r"\cdot",
    "\u25a1": r"\square",
    "\u25cb": r"\circ",
    "\u2299": r"\odot",
}

def _omml_text_to_latex(text: str) -> str:
    """将 OMML <m:t> 文本中的特殊字符映射到 LaTeX。"""
    result = []
    for ch in text:
        result.append(_OMML_SYM_MAP.get(ch, ch))
    return "".join(result)

def _needs_braces(s: str) -> str:
    """如果 s 超过一个 token 需要花括号包裹。"""
    s = s.strip()
    if not s:
        return "{}"
    # 单个字符或单个 \cmd 不需要花括号
    if len(s) == 1:
        return s
    if s.startswith("\\") and " " not in s and "{" not in s:
        return s
    return "{" + s + "}"

def _omml_to_latex(el) -> str:
    """递归将 OMML 元素转为 LaTeX 字符串。"""
    tag = el.tag

    # 纯文本 run
    if tag == _m("t"):
        return _omml_text_to_latex(el.text or "")

    # math run <m:r>: 收集样式+文本
    if tag == _m("r"):
        sty_el = el.find(_m("rPr") + "/" + _m("sty"))  # 注意：find 不支持这种写法
        # 改用直接查找
        rpr = el.find(_m("rPr"))
        sty = None
        if rpr is not None:
            sty_el = rpr.find(_m("sty"))
            if sty_el is not None:
                sty = sty_el.get(_m("val"))
        t_el = el.find(_m("t"))
        text = _omml_text_to_latex(t_el.text or "") if t_el is not None else ""
        if not text:
            return ""
        if sty in ("b", "bi"):
            return r"\mathbf{" + text + "}"
        if sty == "i":
            return r"\mathit{" + text + "}"
        # 普通：单字母默认斜体，多字符用 \mathrm
        if len(text) == 1 and text.isalpha():
            return text  # 默认斜体
        if text.replace(".", "").replace("-", "").replace("+", "").isdigit() or \
           all(c in "0123456789.,-+eE" for c in text):
            return text  # 数字
        return text

    # 上标 <m:sSup>
    if tag == _m("sSup"):
        base = el.find(_m("e"))
        sup = el.find(_m("sup"))
        b = _omml_to_latex(base) if base is not None else ""
        s = _omml_to_latex(sup) if sup is not None else ""
        return f"{_needs_braces(b)}^{_needs_braces(s)}"

    # 下标 <m:sSub>
    if tag == _m("sSub"):
        base = el.find(_m("e"))
        sub = el.find(_m("sub"))
        b = _omml_to_latex(base) if base is not None else ""
        s = _omml_to_latex(sub) if sub is not None else ""
        return f"{_needs_braces(b)}_{_needs_braces(s)}"

    # 上下标 <m:sSubSup>
    if tag == _m("sSubSup"):
        base = el.find(_m("e"))
        sub = el.find(_m("sub"))
        sup = el.find(_m("sup"))
        b = _omml_to_latex(base) if base is not None else ""
        sb = _omml_to_latex(sub) if sub is not None else ""
        sp = _omml_to_latex(sup) if sup is not None else ""
        return f"{_needs_braces(b)}_{_needs_braces(sb)}^{_needs_braces(sp)}"

    # 分数 <m:f>
    if tag == _m("f"):
        num = el.find(_m("num"))
        den = el.find(_m("den"))
        n = _omml_to_latex(num) if num is not None else ""
        d = _omml_to_latex(den) if den is not None else ""
        return r"\frac{" + n + "}{" + d + "}"

    # 根号 <m:rad>
    if tag == _m("rad"):
        deg = el.find(_m("deg"))
        base = el.find(_m("e"))
        b = _omml_to_latex(base) if base is not None else ""
        deg_text = _omml_to_latex(deg).strip() if deg is not None else ""
        if deg_text and deg_text != "":
            return r"\sqrt[" + deg_text + "]{" + b + "}"
        return r"\sqrt{" + b + "}"

    # 极限/求和/积分 <m:nary>
    if tag == _m("nary"):
        pr = el.find(_m("naryPr"))
        chr_el = pr.find(_m("chr")) if pr is not None else None
        chr_val = chr_el.get(_m("val"), "∫") if chr_el is not None else "∫"
        op = _OMML_SYM_MAP.get(chr_val, chr_val)
        sub = el.find(_m("sub"))
        sup = el.find(_m("sup"))
        base = el.find(_m("e"))
        s = _omml_to_latex(sub).strip() if sub is not None else ""
        p = _omml_to_latex(sup).strip() if sup is not None else ""
        b = _omml_to_latex(base) if base is not None else ""
        result = op
        if s:
            result += "_{" + s + "}"
        if p:
            result += "^{" + p + "}"
        result += " " + b
        return result

    # 括号 <m:d>
    if tag == _m("d"):
        pr = el.find(_m("dPr"))
        beg_chr = "("
        end_chr = ")"
        if pr is not None:
            beg_el = pr.find(_m("begChr"))
            end_el = pr.find(_m("endChr"))
            sep_el = pr.find(_m("sepChr"))
            if beg_el is not None:
                beg_chr = beg_el.get(_m("val"), "(")
            if end_el is not None:
                end_chr = end_el.get(_m("val"), ")")
        # 内容
        parts = [_omml_to_latex(e) for e in el.findall(_m("e"))]
        inner = ", ".join(parts)
        # 映射括号字符
        _lmap = {"(": r"\left(", "[": r"\left[", "{": r"\left\{",
                 "|": r"\left|", "‖": r"\left\|",
                 "⌈": r"\left\lceil", "⌊": r"\left\lfloor",
                 "⟨": r"\left\langle", "": "", " ": ""}
        _rmap = {")": r"\right)", "]": r"\right]", "}": r"\right\}",
                 "|": r"\right|", "‖": r"\right\|",
                 "⌉": r"\right\rceil", "⌋": r"\right\rfloor",
                 "⟩": r"\right\rangle", "": "", " ": ""}
        lb = _lmap.get(beg_chr, r"\left" + beg_chr)
        rb = _rmap.get(end_chr, r"\right" + end_chr)
        if lb == "" and rb == "":
            return inner
        return lb + inner + rb

    # 矩阵 <m:m>
    if tag == _m("m"):
        rows = []
        for mr in el.findall(_m("mr")):
            cells = [_omml_to_latex(e) for e in mr.findall(_m("e"))]
            rows.append(" & ".join(cells))
        return r"\begin{pmatrix}" + r" \\ ".join(rows) + r"\end{pmatrix}"

    # 函数 <m:func>
    if tag == _m("func"):
        fname = el.find(_m("fName"))
        base = el.find(_m("e"))
        fn = _omml_to_latex(fname) if fname is not None else ""
        b = _omml_to_latex(base) if base is not None else ""
        return fn + b

    # 累积极限 <m:limLow> / <m:limUpp>
    if tag == _m("limLow"):
        base = el.find(_m("e"))
        lim = el.find(_m("lim"))
        b = _omml_to_latex(base) if base is not None else ""
        l = _omml_to_latex(lim) if lim is not None else ""
        return b + "_{" + l + "}"

    if tag == _m("limUpp"):
        base = el.find(_m("e"))
        lim = el.find(_m("lim"))
        b = _omml_to_latex(base) if base is not None else ""
        l = _omml_to_latex(lim) if lim is not None else ""
        return b + "^{" + l + "}"

    # 过线/下划线 <m:bar>
    if tag == _m("bar"):
        pr = el.find(_m("barPr"))
        pos_el = pr.find(_m("pos")) if pr is not None else None
        pos = pos_el.get(_m("val"), "top") if pos_el is not None else "top"
        base = el.find(_m("e"))
        b = _omml_to_latex(base) if base is not None else ""
        if pos == "bot":
            return r"\underline{" + b + "}"
        return r"\overline{" + b + "}"

    # accent <m:acc>
    if tag == _m("acc"):
        pr = el.find(_m("accPr"))
        chr_el = pr.find(_m("chr")) if pr is not None else None
        acc_ch = chr_el.get(_m("val"), "̂") if chr_el is not None else "̂"
        base = el.find(_m("e"))
        b = _omml_to_latex(base) if base is not None else ""
        _acc_map = {
            "̂": r"\hat", "̃": r"\tilde", "̄": r"\bar", "̇": r"\dot",
            "̈": r"\ddot", "⃗": r"\vec", "̆": r"\breve", "̌": r"\check",
        }
        cmd = _acc_map.get(acc_ch, r"\hat")
        return cmd + "{" + b + "}"

    # eqArr (方程组) <m:eqArr>
    if tag == _m("eqArr"):
        rows = [_omml_to_latex(e) for e in el.findall(_m("e"))]
        return r"\begin{cases}" + r" \\ ".join(rows) + r"\end{cases}"

    # groupChr <m:groupChr>
    if tag == _m("groupChr"):
        pr = el.find(_m("groupChrPr"))
        chr_el = pr.find(_m("chr")) if pr is not None else None
        ch = chr_el.get(_m("val"), "⏞") if chr_el is not None else "⏞"
        pos_el = pr.find(_m("pos")) if pr is not None else None
        pos = pos_el.get(_m("val"), "top") if pos_el is not None else "top"
        base = el.find(_m("e"))
        b = _omml_to_latex(base) if base is not None else ""
        if pos == "bot":
            return r"\underbrace{" + b + "}"
        return r"\overbrace{" + b + "}"

    # phant (phantom)
    if tag == _m("phant"):
        base = el.find(_m("e"))
        return _omml_to_latex(base) if base is not None else ""

    # 通用容器：递归所有子元素拼接
    parts = []
    for child in el:
        # 跳过 ctrlPr (控制属性，不产生输出)
        if child.tag in (_m("ctrlPr"), _m("rPr"), _m("sSubSupPr"), _m("sSupPr"),
                         _m("sSubPr"), _m("fPr"), _m("radPr"), _m("dPr"),
                         _m("mPr"), _m("naryPr"), _m("funcPr"), _m("barPr"),
                         _m("accPr"), _m("groupChrPr"), _m("eqArrPr"),
                         _m("limLowPr"), _m("limUppPr")):
            continue
        parts.append(_omml_to_latex(child))
    return "".join(parts)


def _omml_element_to_latex(omath_el) -> str:
    """将一个 <m:oMath> 元素转为 LaTeX（不含 $ 包裹）。"""
    return _omml_to_latex(omath_el).strip()


def _extract_equations_from_docx(docx_bytes: bytes) -> list[tuple[int, str]]:
    """从 DOCX 中提取公式并转为 LaTeX。

    只遍历 body 的直接子 <w:p>，与 doc.paragraphs 索引对齐。
    返回 (para_idx, latex_str) 列表。
    """
    equations = []
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            if "word/document.xml" not in zf.namelist():
                return equations
            with zf.open("word/document.xml") as f:
                tree = ET.parse(f)

        W = _W_NS_MAIN
        W_BODY = f"{{{W}}}body"
        W_P = f"{{{W}}}p"
        M_OMATH = _m("oMath")

        body = tree.getroot().find(W_BODY)
        if body is None:
            return equations

        for para_idx, p in enumerate(child for child in body if child.tag == W_P):
            latex_parts = []
            for omath in p.iter(M_OMATH):
                latex = _omml_element_to_latex(omath)
                if latex:
                    latex_parts.append(latex)
            if latex_parts:
                equations.append((para_idx, r" \quad ".join(latex_parts)))
    except Exception:
        pass
    return equations



def _build_original_html(docx_bytes: bytes) -> str:
    """将 DOCX 转换为 HTML 用于原文预览。

    在每个顶层 <p> 上注入 data-para-idx 属性，供格式化预览克隆后按索引查找段落。
    同时将 OMML 公式转换为 LaTeX 注入对应段落，供 KaTeX 渲染。
    """
    import html as html_lib

    style_map = """
    p[style-name='Heading 1'] => h1:fresh
    p[style-name='Heading 2'] => h2:fresh
    """
    try:
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes), style_map=style_map)
        html_str = result.value

        # 构建 body 直接子元素类型序列（p/tbl），用于对齐 HTML 输出
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W_BODY = f"{{{W_NS}}}body"
        W_P    = f"{{{W_NS}}}p"
        W_TBL  = f"{{{W_NS}}}tbl"
        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                with zf.open("word/document.xml") as xf:
                    xml_tree = ET.parse(xf)
            body_el = xml_tree.getroot().find(W_BODY)
            body_seq = [
                "p" if c.tag == W_P else "tbl"
                for c in body_el
                if c.tag in (W_P, W_TBL)
            ]
        except Exception:
            body_seq = []

        # 公式映射：para_idx -> latex
        equations = _extract_equations_from_docx(docx_bytes)
        eq_map = {idx: formula for idx, formula in equations}

        if body_seq:
            # 将 HTML 拆分为顶层块（<p...>...</p> 和 <table>...</table>）
            parts = re.split(
                r'(<p[^>]*>.*?</p>|<table\b[^>]*>.*?</table>)',
                html_str, flags=re.DOTALL
            )
            # 找出 parts 中每个块的索引
            block_part_indices = [
                i for i, s in enumerate(parts)
                if re.match(r'<p\b|<table\b', s)
            ]

            # p 的计数器（para_idx 只计 w:p，不计 w:tbl）
            para_idx = 0
            for seq_idx, block_type in enumerate(body_seq):
                if seq_idx >= len(block_part_indices):
                    break
                pi = block_part_indices[seq_idx]

                if block_type == "p":
                    p_html = parts[pi]
                    # 在 <p 后注入 data-para-idx
                    p_html = re.sub(
                        r'^(<p\b)',
                        rf'\1 data-para-idx="{para_idx}"',
                        p_html
                    )
                    # 注入公式（如果该段有公式）
                    if para_idx in eq_map:
                        formula = eq_map[para_idx]
                        eq_span = (
                            f' <span class="katex-formula" data-latex="{html_lib.escape(formula)}"'
                            f' style="font-family:\'Times New Roman\',serif;font-size:10.5pt;">'
                            f'\\({html_lib.escape(formula)}\\)</span>'
                        )
                        p_html = p_html[:-4] + eq_span + "</p>"
                    parts[pi] = p_html
                    para_idx += 1

            html_str = "".join(parts)

        widths = _extract_image_widths_cm(docx_bytes)
        html_str = _inject_image_widths(html_str, widths)
        anchors = _extract_media_anchors(docx_bytes)
        return _inject_media_ids(html_str, anchors)
    except Exception as e:
        import logging
        logging.error(f"Mammoth conversion error: {e}")
        return "<p>文档预览生成失败</p>"


def _parse_docx(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    ctx = ChapterContext()
    paragraphs = []
    
    # 提取公式信息
    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()
    equations = _extract_equations_from_docx(docx_bytes)
    equation_map = {idx: formula for idx, formula in equations}
    
    for i, para in enumerate(doc.paragraphs):
        result = classify_paragraph(para, ctx)
        result["index"] = i
        result["is_formula"] = i in equation_map
        result["formula_text"] = equation_map.get(i, None)

        if i in equation_map:
            formula = equation_map[i]
            existing_text = result.get("original_text", "").strip()
            if existing_text:
                # 段落同时包含普通文本和公式，将公式文本追加到末尾
                result["original_text"] = f"{existing_text} [{formula}]"
                # 保持原始 detected_level，不强制改为 Equation
            else:
                # 纯公式段落（doc.text 为空）
                result["original_text"] = formula
                result["detected_level"] = "Equation"

        if result["detected_level"] != "Body" and result["detected_level"] != "Equation":
            ctx.push(result["detected_level"], para.text[:20])
        paragraphs.append(result)
    return paragraphs


@app.get("/templates")
def get_templates():
    return list_templates(TEMPLATES_DIR)


@app.get("/template")
def get_template(name: str = "report_cn.yaml"):
    return load_template(os.path.join(TEMPLATES_DIR, name))


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    # 验证文件扩展名
    filename = file.filename.lower()
    if not (filename.endswith(".doc") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="仅支持 .doc 和 .docx 格式的文件")
    
    doc_id = uuid.uuid4().hex
    sess = _session_dir(doc_id)
    os.makedirs(sess, exist_ok=True)

    raw_path = os.path.join(sess, "original_" + file.filename)
    with open(raw_path, "wb") as f:
        f.write(await file.read())

    was_doc = is_doc(raw_path)
    try:
        if was_doc:
            docx_path = to_docx(raw_path, sess)
        else:
            docx_path = raw_path

        paragraphs = _parse_docx(docx_path)
    except Exception as e:
        # 清理已创建的会话目录
        shutil.rmtree(sess, ignore_errors=True)
        raise HTTPException(status_code=400, detail=f"无法解析文档: {str(e)}")

    with open(docx_path, "rb") as f:
        docx_bytes = f.read()
    media_anchors = _extract_media_anchors(docx_bytes)
    original_html = _build_original_html(docx_bytes)

    meta = {
        "doc_id": doc_id,
        "filename": file.filename,
        "was_doc": was_doc,
        "raw_path": raw_path,
        "docx_path": docx_path,
        "paragraphs": paragraphs,
        "media_anchors": media_anchors,
        "original_html": original_html,
    }
    with open(os.path.join(sess, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False)

    return {
        "doc_id": doc_id,
        "filename": file.filename,
        "was_doc": was_doc,
        "paragraphs": paragraphs,
        "media_anchors": media_anchors,
        "original_html": original_html,
    }


@app.get("/session/{doc_id}")
def get_session(doc_id: str):
    sess = _session_dir(doc_id)
    meta_path = os.path.join(sess, "meta.json")
    if not os.path.isfile(meta_path):
        raise HTTPException(status_code=404, detail="session not found or expired")
    _touch(sess)
    _touch(meta_path)
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    return {
        "doc_id": meta["doc_id"],
        "filename": meta["filename"],
        "was_doc": meta["was_doc"],
        "paragraphs": meta["paragraphs"],
        "media_anchors": meta["media_anchors"],
        "original_html": meta["original_html"],
    }


@app.delete("/session/{doc_id}")
def delete_session(doc_id: str):
    sess = _session_dir(doc_id)
    shutil.rmtree(sess, ignore_errors=True)
    return {"ok": True}


@app.post("/original-preview")
async def original_preview(file: UploadFile = File(...)):
    content = await file.read()
    return {"html": _build_original_html(content)}


class PreviewRequest(BaseModel):
    text: str
    level: str
    template: str = "report_cn.yaml"


@app.post("/preview")
def preview(req: PreviewRequest):
    tpl = load_template(os.path.join(TEMPLATES_DIR, req.template))
    html = render_paragraph_html(req.text, req.level, tpl)
    return {"html": html}


class PreviewBatchItem(BaseModel):
    text: str
    level: str


class PreviewBatchRequest(BaseModel):
    items: list[PreviewBatchItem]
    template: str = "report_cn.yaml"


def _build_styled_html(docx_bytes: bytes, template_name: str) -> str:
    """
    生成格式化预览 HTML：
    - 遍历 DOCX body 直接子元素（w:p 和 w:tbl）
    - w:p 段落：用模板样式渲染，保留图片（<img> 从原文 HTML 中提取）
    - w:tbl 表格：直接使用 mammoth 生成的 <table> HTML
    - 公式：注入 KaTeX LaTeX span
    """
    import html as html_lib

    tpl = load_template(os.path.join(TEMPLATES_DIR, template_name))

    # 1. mammoth 生成原文 HTML，提取表格和图片
    style_map = """
    p[style-name='Heading 1'] => h1:fresh
    p[style-name='Heading 2'] => h2:fresh
    """
    try:
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes), style_map=style_map)
        orig_html = result.value
    except Exception:
        orig_html = ""

    # 从 mammoth HTML 提取顶层 <table> 块（按顺序）
    tbl_parts = re.findall(r'<table\b[^>]*>.*?</table>', orig_html, flags=re.DOTALL)

    # 从 mammoth HTML 提取所有 <img> 标签（base64），按出现顺序
    img_tags = re.findall(r'<img[^>]+>', orig_html)
    img_idx = 0

    # 2. 解析 DOCX XML
    W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_BODY = f"{{{W_NS}}}body"
    W_P    = f"{{{W_NS}}}p"
    W_TBL  = f"{{{W_NS}}}tbl"
    M_NS   = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    M_OMATH = f"{{{M_NS}}}oMath"
    _WP_DRAW = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            with zf.open("word/document.xml") as xf:
                xml_tree = ET.parse(xf)
        body_el = xml_tree.getroot().find(W_BODY)
    except Exception:
        return "<p>格式化预览生成失败</p>"

    # 3. 公式映射
    equations = _extract_equations_from_docx(docx_bytes)
    eq_map = {idx: formula for idx, formula in equations}

    # 4. 遍历并生成 HTML
    doc = Document(io.BytesIO(docx_bytes))
    ctx = ChapterContext()
    parts_html = []
    para_idx = 0
    tbl_idx = 0

    for child in body_el:
        if child.tag == W_P:
            # 找到对应的 doc.paragraphs[para_idx]
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
            else:
                para_idx += 1
                continue

            result = classify_paragraph(para, ctx)
            level = result.get("detected_level", "Body")
            text  = result.get("original_text", "") or ""

            if level not in ("Body", "Equation"):
                ctx.push(level, text[:20])

            # 公式处理
            if para_idx in eq_map:
                formula = eq_map[para_idx]
                if not text.strip():
                    level = "Equation"
                    text  = formula

            # 检查段落是否含图片（<wp:inline> or <wp:anchor>）
            has_image = child.find(f".//{_WP_DRAW}inline") is not None or \
                        child.find(f".//{_WP_DRAW}anchor") is not None

            if has_image and img_idx < len(img_tags):
                # 含图片的段落：插入 <img>，不改图片样式
                img_html = img_tags[img_idx]
                img_idx += 1
                eq_span = ""
                if para_idx in eq_map:
                    formula = eq_map[para_idx]
                    escaped = html_lib.escape(formula)
                    eq_span = (
                        f' <span class="katex-formula" data-latex="{escaped}"'
                        f' style="font-family:\'Times New Roman\',serif;font-size:10.5pt;">'
                        f'\\({escaped}\\)</span>'
                    )
                parts_html.append(
                    f'<p data-para-idx="{para_idx}" style="text-align:center;margin:4px 0">'
                    f'{img_html}{eq_span}</p>'
                )
            else:
                # 普通段落：应用模板样式
                p_html = render_paragraph_html(text, level, tpl)
                # 注入 data-para-idx
                p_html = re.sub(r'^(<p\b)', rf'\1 data-para-idx="{para_idx}"', p_html)

                # 如果有公式且段落非纯公式段，在末尾追加公式 span
                if para_idx in eq_map and level != "Equation":
                    formula = eq_map[para_idx]
                    escaped = html_lib.escape(formula)
                    eq_span = (
                        f' <span class="katex-formula" data-latex="{escaped}"'
                        f' style="font-family:\'Times New Roman\',serif;font-size:10.5pt;">'
                        f'\\({escaped}\\)</span>'
                    )
                    p_html = p_html[:-4] + eq_span + "</p>"

                parts_html.append(p_html)

            para_idx += 1

        elif child.tag == W_TBL:
            if tbl_idx < len(tbl_parts):
                # 使用 mammoth 生成的表格 HTML
                parts_html.append(
                    f'<div style="overflow-x:auto;margin:8px 0">{tbl_parts[tbl_idx]}</div>'
                )
                tbl_idx += 1
            else:
                parts_html.append('<div style="color:#999;text-align:center;padding:8px">[表格]</div>')

    return "\n".join(parts_html)


@app.post("/preview/full-html")
def preview_full_html(doc_id: str = Form(...), template: str = Form("report_cn.yaml")):
    sess = _session_dir(doc_id)
    meta_path = os.path.join(sess, "meta.json")
    if not os.path.isfile(meta_path):
        raise HTTPException(status_code=404, detail="session not found")
    _touch(sess)
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)
    with open(meta["docx_path"], "rb") as f:
        docx_bytes = f.read()
    styled_html = _build_styled_html(docx_bytes, template)
    return {"html": styled_html}


@app.post("/preview/batch")
def preview_batch(req: PreviewBatchRequest):
    tpl = load_template(os.path.join(TEMPLATES_DIR, req.template))
    return {"htmls": [render_paragraph_html(it.text, it.level, tpl) for it in req.items]}


def _replace_para_text(para: Paragraph, new_text: str):
    """把段落的所有 run 合并成一个，文本替换为 new_text，保留第一个 run 的 rPr。"""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _insert_para_after(para: Paragraph, text: str) -> Paragraph:
    """在 para 之后插入一个新 w:p，复制 para 的 pPr，文本为 text。返回新 Paragraph。"""
    new_p = OxmlElement("w:p")
    pPr = para._element.find(qn_w("pPr"))
    if pPr is not None:
        from copy import deepcopy
        new_p.append(deepcopy(pPr))
    para._element.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    new_para.add_run(text)
    return new_para


def qn_w(tag: str) -> str:
    return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + tag


@app.post("/export")
async def export_doc(
    background_tasks: BackgroundTasks,
    paragraphs: str = Form(...),
    template: str = Form("report_cn.yaml"),
    doc_id: str = Form(...),
):
    items = json.loads(paragraphs)
    tpl = load_template(os.path.join(TEMPLATES_DIR, template))

    sess = _session_dir(doc_id)
    meta_path = os.path.join(sess, "meta.json")
    if not os.path.isfile(meta_path):
        raise HTTPException(status_code=404, detail="session not found or expired")
    _touch(sess)
    _touch(meta_path)
    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    src_docx = meta["docx_path"]
    filename = meta["filename"]
    was_doc = meta["was_doc"]

    out_dir = tempfile.mkdtemp()
    background_tasks.add_task(shutil.rmtree, out_dir, True)

    base = os.path.splitext(filename)[0]
    out_docx = os.path.join(out_dir, f"{base}_formatted.docx")
    shutil.copy2(src_docx, out_docx)

    doc = Document(out_docx)
    apply_page_margins(doc, tpl)

    para_list = list(doc.paragraphs)
    by_index = {int(it["index"]): it for it in items if "index" in it}

    for idx, para in enumerate(para_list):
        item = by_index.get(idx)
        if item is None:
            continue
        level = item.get("detected_level", "Body")

        if item.get("is_split") and item.get("split_title") and item.get("split_body"):
            _replace_para_text(para, item["split_title"])
            apply_format_to_paragraph(para, level, tpl)
            body_para = _insert_para_after(para, item["split_body"])
            apply_format_to_paragraph(body_para, "Body", tpl)
        else:
            apply_format_to_paragraph(para, level, tpl)

    center_tables_and_images(doc)

    doc.save(out_docx)

    if was_doc:
        out_path = to_doc(out_docx, out_dir)
        media_type = "application/msword"
    else:
        out_path = out_docx
        media_type = ("application/vnd.openxmlformats-officedocument"
                      ".wordprocessingml.document")

    return FileResponse(out_path, media_type=media_type,
                        filename=os.path.basename(out_path))
