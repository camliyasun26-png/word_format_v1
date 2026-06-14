# formatter/applier.py
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ALIGNMENT_MAP = {
    "CENTER":  WD_ALIGN_PARAGRAPH.CENTER,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "LEFT":    WD_ALIGN_PARAGRAPH.LEFT,
    "RIGHT":   WD_ALIGN_PARAGRAPH.RIGHT,
}

_LEVEL_ORDER = ["h1", "h2", "h3", "h4", "h5"]


def _max_preset_level(template: dict) -> str:
    for lvl in reversed(_LEVEL_ORDER):
        if lvl in template:
            return lvl
    return "body"


def _set_run_font_all_slots(run, font_name: str):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    run.font.name = font_name


def _disable_snap_to_grid(para):
    pPr = para._element.get_or_add_pPr()
    snap = pPr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        pPr.append(snap)
    snap.set(qn("w:val"), "0")


def _set_outline_level(para, level_num):
    pPr = para._element.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if level_num is None:
        if ol is not None:
            pPr.remove(ol)
        return
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        pPr.append(ol)
    ol.set(qn("w:val"), str(level_num))


def apply_format_to_paragraph(para, level: str, template: dict):
    key = level.lower() if level != "Body" else "body"

    if key.startswith("h") and key[1:].isdigit():
        _set_outline_level(para, int(key[1:]) - 1)
    else:
        _set_outline_level(para, None)

    if key not in template:
        hn_key = _max_preset_level(template)
        font_name = template[hn_key]["font_name"]
        for run in para.runs:
            _set_run_font_all_slots(run, font_name)
        return

    cfg = template[key]
    font_name    = cfg["font_name"]
    font_size    = Pt(cfg["font_size"])
    bold         = cfg.get("bold")
    alignment    = _ALIGNMENT_MAP.get(cfg.get("alignment", "JUSTIFY"),
                                      WD_ALIGN_PARAGRAPH.JUSTIFY)
    indent_pt    = cfg.get("indent_first_line", 0)
    line_spacing = cfg.get("line_spacing", 1.5)
    space_before = Pt(cfg.get("space_before", 0))
    space_after  = Pt(cfg.get("space_after", 0))

    para.alignment = alignment
    pf = para.paragraph_format
    pf.first_line_indent = Pt(indent_pt)
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after  = space_after
    _disable_snap_to_grid(para)

    for run in para.runs:
        _set_run_font_all_slots(run, font_name)
        run.font.size = font_size
        if bold is True:
            run.bold = True
        elif bold is False:
            run.bold = False


def apply_font_only(para, font_name: str):
    for run in para.runs:
        _set_run_font_all_slots(run, font_name)


def center_tables_and_images(doc: Document):
    """对表格和图片段落施加居中。原本就居中的不变（写入等价值，不影响其他格式）。"""
    for tbl in doc.tables:
        tblPr = tbl._element.find(qn("w:tblPr"))
        if tblPr is not None:
            tblInd = tblPr.find(qn("w:tblInd"))
            if tblInd is not None:
                tblPr.remove(tblInd)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._element.find(qn("w:tblPr"))
        if tblPr is not None:
            jc = tblPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                tblPr.append(jc)
            jc.set(qn("w:val"), "center")

    drawing_tag = qn("w:drawing")
    pict_tag = qn("w:pict")
    for para in doc.paragraphs:
        has_media = False
        for el in para._element.iter():
            if el.tag == drawing_tag or el.tag == pict_tag:
                has_media = True
                break
        if has_media:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None:
                    pPr.remove(ind)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_page_margins(doc: Document, template: dict):
    """应用页面边距设置，支持装订线、页眉/页脚距离"""
    margins = template.get("page_margins_cm", {})
    for section in doc.sections:
        # 标准边距
        if "top" in margins:
            section.top_margin = Cm(margins["top"])
        if "bottom" in margins:
            section.bottom_margin = Cm(margins["bottom"])
        if "left" in margins:
            section.left_margin = Cm(margins["left"])
        if "right" in margins:
            section.right_margin = Cm(margins["right"])
        
        # 装订线 (gutter) - 兼容处理
        gutter = margins.get("gutter")
        if gutter is not None:
            section.gutter = Cm(gutter)
        
        # 页眉距离 (header_distance_cm) - 兼容处理
        header_distance = template.get("header_distance_cm")
        if header_distance is not None:
            section.header_distance = Cm(header_distance)
        
        # 页脚距离 (footer_distance_cm) - 兼容处理
        footer_distance = template.get("footer_distance_cm")
        if footer_distance is not None:
            section.footer_distance = Cm(footer_distance)


def apply_header_settings(doc: Document, template: dict):
    """应用页眉设置"""
    header_config = template.get("header", {})
    
    # 如果未启用或未配置，跳过
    if not header_config.get("enabled", False):
        for section in doc.sections:
            # 清除页眉
            header = section.header
            for para in header.paragraphs:
                para.clear()
        return
    
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # 清除现有内容
        for para in header.paragraphs:
            para.clear()
        
        # 设置页眉内容
        content = header_config.get("content_left", "")
        if not content:
            content = header_config.get("content_center", "")
        
        if content:
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.text = content
            
            # 设置字体
            font_name = header_config.get("font_name", "宋体")
            font_size = header_config.get("font_size", 9.0)
            para.runs[0].font.name = font_name if header.paragraphs else None
            if para.runs:
                para.runs[0].font.size = Pt(font_size)
                _set_run_font_all_slots(para.runs[0], font_name)
            
            # 设置对齐方式
            alignment = header_config.get("alignment", "LEFT")
            para.alignment = _ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 页眉下划线 - 兼容处理
        if header_config.get("underline", False):
            for para in header.paragraphs:
                pPr = para._element.get_or_add_pPr()
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is None:
                    pBdr = OxmlElement("w:pBdr")
                    pPr.append(pBdr)
                bottom = pBdr.find(qn("w:bottom"))
                if bottom is None:
                    bottom = OxmlElement("w:bottom")
                    pBdr.append(bottom)
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "auto")


def apply_footer_settings(doc: Document, template: dict):
    """应用页脚设置"""
    footer_config = template.get("footer", {})
    
    # 如果未启用或未配置，跳过
    if not footer_config.get("enabled", False):
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                para.clear()
        return
    
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 清除现有内容
        for para in footer.paragraphs:
            para.clear()
        
        # 页码位置
        page_position = footer_config.get("page_number_position", "CENTER")
        
        # Word会自动处理页码，这里可以添加页脚文字
        content = footer_config.get("content", "")
        if content:
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.text = content
            
            font_name = footer_config.get("font_name", "宋体")
            font_size = footer_config.get("font_size", 9.0)
            if para.runs:
                para.runs[0].font.size = Pt(font_size)
                _set_run_font_all_slots(para.runs[0], font_name)
            
            para.alignment = _ALIGNMENT_MAP.get(page_position, WD_ALIGN_PARAGRAPH.CENTER)


def apply_page_break_before(para, template: dict, level: str):
    """在指定级别的段落前添加分页符"""
    key = level.lower() if level != "Body" else "body"
    
    # 获取该级别的page_break_before设置
    level_config = template.get(key, {})
    if not level_config.get("page_break_before", False):
        return
    
    # 添加分页符
    pPr = para._element.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    page_break.set(qn("w:val"), "1")
    pPr.append(page_break)


def apply_figure_format(paragraphs: list, template: dict):
    """应用图片格式设置（图片标题）"""
    figure_config = template.get("figure", {})
    if not figure_config:
        return
    
    caption_config = figure_config.get("caption", {})
    if not caption_config:
        return
    
    font_name = caption_config.get("font_name", "宋体")
    font_size = Pt(caption_config.get("font_size", 9.0))
    alignment = _ALIGNMENT_MAP.get(caption_config.get("alignment", "CENTER"), 
                                    WD_ALIGN_PARAGRAPH.CENTER)
    
    # 识别图片标题（包含"图"字的段落）
    for para in paragraphs:
        text = para.text.strip()
        if text.startswith("图") and "：" in text or text.startswith("图") and ":" in text:
            para.alignment = alignment
            for run in para.runs:
                run.font.size = font_size
                _set_run_font_all_slots(run, font_name)


def apply_table_format(doc: Document, template: dict):
    """应用表格格式设置"""
    table_config = template.get("table", {})
    if not table_config:
        return
    
    caption_config = table_config.get("caption", {})
    if not caption_config:
        return
    
    font_name = caption_config.get("font_name", "宋体")
    font_size = Pt(caption_config.get("font_size", 9.0))
    alignment = _ALIGNMENT_MAP.get(caption_config.get("alignment", "CENTER"),
                                    WD_ALIGN_PARAGRAPH.CENTER)
    
    # 识别表格标题（包含"表"字的段落）
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("表") and "：" in text or text.startswith("表") and ":" in text:
            para.alignment = alignment
            for run in para.runs:
                run.font.size = font_size
                _set_run_font_all_slots(run, font_name)
    
    # 三线表样式 - 兼容处理
    table_style = table_config.get("style")
    if table_style == "THREE_LINE":
        for tbl in doc.tables:
            _apply_three_line_table_style(tbl)


def _apply_three_line_table_style(tbl):
    """应用三线表样式"""
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._element.insert(0, tblPr)
    
    # 设置表格宽度
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    
    # 设置边框
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    
    # 上边框（粗线）
    for border_name in ["top"]:
        border = tblBorders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tblBorders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")  # 1.5pt
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
    
    # 下边框（粗线）
    for border_name in ["bottom"]:
        border = tblBorders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tblBorders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")  # 1.5pt
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
    
    # 内部边框（细线或无）
    for border_name in ["left", "right", "insideH", "insideV"]:
        border = tblBorders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tblBorders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # 0.5pt
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def apply_footnote_format(doc: Document, template: dict):
    """应用脚注格式设置"""
    footnote_config = template.get("footnote", {})
    if not footnote_config:
        return
    
    font_name = footnote_config.get("font_name", "宋体")
    font_size = Pt(footnote_config.get("font_size", 9.0))
    
    # 获取脚注引用符的格式
    footnotes_part = None
    try:
        footnotes_part = doc.part.package.part_related_by("http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    except Exception:
        pass
    
    if footnotes_part is None:
        return
    
    # 设置脚注分隔线
    if footnote_config.get("separator_line", False):
        for section in doc.sections:
            footer = section.footer
            pPr = footer._element.get_or_add_pPr()
            
            # 查找或创建分隔线
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = OxmlElement("w:pBdr")
                pPr.append(pBdr)
            
            bottom = pBdr.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                pBdr.append(bottom)
            
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")


def apply_special_paragraph_format(doc: Document, template: dict):
    """应用特殊段落格式（摘要、关键词、参考文献、致谢等）"""
    
    # 摘要格式
    abstract_config = template.get("abstract", {})
    if abstract_config:
        _apply_paragraph_format_by_keyword(doc, "摘要", abstract_config)
    
    # 关键词格式
    keywords_config = template.get("keywords", {})
    if keywords_config:
        _apply_paragraph_format_by_keyword(doc, "关键词", keywords_config)
    
    # 参考文献格式
    reference_config = template.get("reference", {})
    if reference_config:
        _apply_paragraph_format_by_keyword(doc, "参考文献", reference_config)
    
    # 致谢格式
    acknowledgment_config = template.get("acknowledgment", {})
    if acknowledgment_config:
        _apply_paragraph_format_by_keyword(doc, "致谢", acknowledgment_config)
    
    # 公式格式
    equation_config = template.get("equation", {})
    if equation_config:
        _apply_equation_format(doc, equation_config)


def _apply_paragraph_format_by_keyword(doc: Document, keyword: str, config: dict):
    """根据关键词匹配应用段落格式"""
    font_name = config.get("font_name", "宋体")
    font_size = Pt(config.get("font_size", 10.5))
    alignment = _ALIGNMENT_MAP.get(config.get("alignment", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT)
    line_spacing = config.get("line_spacing", 1.5)
    indent_first_line = config.get("indent_first_line", 0)
    
    found_keyword = False
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 匹配关键词行
        if keyword in text and not found_keyword:
            found_keyword = True
            para.alignment = alignment
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            if indent_first_line > 0:
                pf.first_line_indent = Pt(indent_first_line)
            
            for run in para.runs:
                run.font.size = font_size
                _set_run_font_all_slots(run, font_name)
        
        # 应用格式到后续段落直到下一个标题
        elif found_keyword:
            # 如果遇到一级标题，停止应用
            if text.startswith("第") and "章" in text:
                break
            # 应用格式
            para.alignment = alignment
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            if indent_first_line > 0:
                pf.first_line_indent = Pt(indent_first_line)
            
            for run in para.runs:
                run.font.size = font_size
                _set_run_font_all_slots(run, font_name)


def _apply_equation_format(doc: Document, config: dict):
    """应用数学公式格式（支持 WPS 公式）"""
    font_name = config.get("font_name", "Times New Roman")
    font_size = Pt(config.get("font_size", 10.5))
    alignment = _ALIGNMENT_MAP.get(config.get("alignment", "RIGHT"), WD_ALIGN_PARAGRAPH.RIGHT)
    
    math_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 检查段落是否包含 MathML 公式（WPS 公式）
        has_math = para._element.find('.//{' + math_ns + '}oMath') is not None
        
        # 识别公式段落：包含 MathML 公式 或 包含等号且较短的段落
        if has_math or ("=" in text and len(text) < 200 and not para.paragraph_format.left_indent):
            para.alignment = alignment
            for run in para.runs:
                run.font.size = font_size
                _set_run_font_all_slots(run, font_name)


def apply_template_full(doc: Document, template: dict):
    """应用完整模板到文档（兼容入口函数）"""
    # 1. 应用页面设置
    apply_page_margins(doc, template)
    
    # 2. 应用页眉页脚
    apply_header_settings(doc, template)
    apply_footer_settings(doc, template)
    
    # 3. 应用段落格式
    for para in doc.paragraphs:
        level = "body"  # 默认使用body级别
        
        # 根据段落样式判断级别
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            outlineLvl = pPr.find(qn("w:outlineLvl"))
            if outlineLvl is not None:
                lvl_val = outlineLvl.get(qn("w:val"))
                if lvl_val is not None:
                    level = f"h{int(lvl_val) + 1}"
        
        # 应用格式
        apply_format_to_paragraph(para, level, template)
        
        # 检查是否需要分页
        apply_page_break_before(para, template, level)
    
    # 4. 应用特殊段落格式
    apply_special_paragraph_format(doc, template)
    
    # 5. 居中表格和图片
    center_tables_and_images(doc)
    
    # 6. 应用图表格式
    apply_table_format(doc, template)
    apply_figure_format(doc.paragraphs, template)
    
    # 7. 应用脚注格式
    apply_footnote_format(doc, template)
