# tests/test_api.py
import pytest
from fastapi.testclient import TestClient

def test_list_templates():
    from main import app
    client = TestClient(app)
    resp = client.get("/templates")
    assert resp.status_code == 200
    data = resp.json()
    assert isinstance(data, list)
    assert any("report_cn" in t for t in data)

def test_upload_docx(tmp_path):
    from main import app
    client = TestClient(app)
    from docx import Document
    doc = Document()
    p = doc.add_paragraph("第一章 总体情况")
    for run in p.runs:
        run.bold = True
    path = tmp_path / "test.docx"
    doc.save(str(path))
    with open(path, "rb") as f:
        resp = client.post("/upload", files={"file": ("test.docx", f,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert resp.status_code == 200
    data = resp.json()
    assert "paragraphs" in data
    assert len(data["paragraphs"]) > 0

def test_preview_paragraph():
    from main import app
    client = TestClient(app)
    resp = client.post("/preview", json={
        "text": "第一章 总体情况",
        "level": "H1",
        "template": "report_cn.yaml"
    })
    assert resp.status_code == 200
    assert "<p " in resp.json()["html"]


def test_original_preview_injects_image_width(tmp_path):
    from main import app
    from docx import Document
    from docx.shared import Cm
    from PIL import Image
    import io as _io

    img = Image.new("RGB", (100, 50), "red")
    buf = _io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)

    doc = Document()
    doc.add_paragraph("hello")
    doc.add_picture(buf, width=Cm(6))
    path = tmp_path / "img.docx"
    doc.save(str(path))

    client = TestClient(app)
    with open(path, "rb") as f:
        resp = client.post("/original-preview", files={"file": ("img.docx", f,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert resp.status_code == 200
    html = resp.json()["html"]
    assert "width:6.00cm" in html
    assert "max-width:100%" in html
