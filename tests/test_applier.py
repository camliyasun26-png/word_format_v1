# tests/test_applier.py
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from formatter.applier import apply_format_to_paragraph, apply_page_margins
from formatter.template_loader import load_template

TPL = load_template("templates/report_cn.yaml")

def _make_doc_para(text: str):
    doc = Document()
    p = doc.add_paragraph(text)
    return doc, p

def test_h1_bold_center():
    doc, p = _make_doc_para("第一章")
    apply_format_to_paragraph(p, "H1", TPL)
    assert p.runs[0].bold is True
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER

def test_body_not_bold():
    doc, p = _make_doc_para("正文")
    apply_format_to_paragraph(p, "Body", TPL)
    assert p.runs[0].bold is False

def test_h2_font_size():
    doc, p = _make_doc_para("一、")
    apply_format_to_paragraph(p, "H2", TPL)
    assert p.runs[0].font.size == Pt(14)

def test_page_margins():
    doc = Document()
    apply_page_margins(doc, TPL)
    section = doc.sections[0]
    assert abs(section.top_margin.cm - 2.54) < 0.01
    assert abs(section.left_margin.cm - 3.17) < 0.01
